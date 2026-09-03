"""
Document Processing & Text Extraction Module.
Extracts clean, structured text and page-level metadata from PDF, DOCX, and TXT files.
Designed as the foundational ingestion layer for the Virtual Hawks RAG pipeline.
"""

from __future__ import annotations

import io
import os
import re
import unicodedata
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, BinaryIO, Dict, List, Optional, Union
import zipfile

import pypdf
import pypdf.errors
import docx
import docx.opc.exceptions

from app.rag.exceptions import (
    CorruptedFileError,
    DocumentProcessingError,
    EmptyDocumentError,
    FileAccessError,
    UnsupportedFileTypeError,
)


@dataclass
class DocumentPage:
    """Represents an extracted page or logical section with text and metadata."""
    page_number: int
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert page object to JSON-serializable dictionary."""
        return {
            "page_number": self.page_number,
            "text": self.text,
            "metadata": self.metadata,
        }


@dataclass
class ProcessedDocument:
    """Standardized representation of an extracted document."""
    file_name: Optional[str]
    file_type: str
    total_pages: int
    pages: List[DocumentPage]
    raw_text: str
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert document object to JSON-serializable dictionary."""
        return {
            "file_name": self.file_name,
            "file_type": self.file_type,
            "total_pages": self.total_pages,
            "pages": [page.to_dict() for page in self.pages],
            "raw_text": self.raw_text,
            "metadata": self.metadata,
        }

    def chunk(self, chunker: Optional[Any] = None, **chunker_kwargs) -> Any:
        """
        Directly chunk this processed document using SemanticChunker (Milestones 1 & 2 bridge).
        """
        from app.rag.chunker import SemanticChunker
        c = chunker or SemanticChunker(**chunker_kwargs)
        return c.chunk_document(self)


class DocumentProcessor:
    """
    Modular, pure-Python utility to extract clean text from standard educational files
    (PDF, DOCX, TXT, MD).
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
    MIME_TYPE_MAPPINGS = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/msword": ".docx",
        "text/plain": ".txt",
        "text/markdown": ".md",
        "text/x-markdown": ".md",
    }

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Normalize Unicode, remove non-printable control characters,
        and trim excessive trailing/leading whitespace while preserving paragraph structure.
        """
        if not text:
            return ""

        normalized = unicodedata.normalize("NFKC", text)
        normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")

        cleaned_chars = [
            char for char in normalized
            if char in ("\n", "\t") or (unicodedata.category(char) != "Cc" and ord(char) >= 32)
        ]
        cleaned_text = "".join(cleaned_chars)

        cleaned_lines = [re.sub(r"[ \t]+$", "", line) for line in cleaned_text.split("\n")]
        cleaned_text = "\n".join(cleaned_lines)

        cleaned_text = re.sub(r"\n{3,}", "\n\n", cleaned_text)
        return cleaned_text.strip()

    @classmethod
    def _resolve_file_type(
        cls,
        file_source: Union[str, Path, BinaryIO, bytes],
        file_type: Optional[str] = None,
        file_name: Optional[str] = None,
    ) -> str:
        """Determines the file extension/type from parameters or file source."""
        ext = None

        if file_type:
            ft_clean = file_type.lower().strip()
            if ft_clean in cls.MIME_TYPE_MAPPINGS:
                ext = cls.MIME_TYPE_MAPPINGS[ft_clean]
            elif ft_clean.startswith("."):
                ext = ft_clean
            else:
                ext = f".{ft_clean}"

        if not ext and file_name:
            ext = Path(file_name).suffix.lower()

        if not ext and isinstance(file_source, (str, Path)):
            ext = Path(file_source).suffix.lower()

        if not ext or ext not in cls.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"Unsupported or indeterminate file type: '{ext or file_type}'. "
                f"Supported types: {', '.join(sorted(cls.SUPPORTED_EXTENSIONS))}"
            )

        return ext

    @classmethod
    def _read_bytes(
        cls,
        file_source: Union[str, Path, BinaryIO, bytes],
    ) -> bytes:
        """Reads raw bytes from various input sources with defensive stream position handling."""
        if isinstance(file_source, bytes):
            return file_source

        if isinstance(file_source, io.BytesIO):
            val = file_source.getvalue()
            return val

        if hasattr(file_source, "read"):
            try:
                if hasattr(file_source, "tell") and hasattr(file_source, "seek"):
                    pos = file_source.tell()
                    content = file_source.read()
                    if not content and pos != 0:
                        file_source.seek(0)
                        content = file_source.read()
                else:
                    content = file_source.read()

                if isinstance(content, str):
                    return content.encode("utf-8")
                return content
            except Exception as e:
                raise FileAccessError(f"Failed to read from file stream: {str(e)}") from e

        if isinstance(file_source, (str, Path)):
            path = Path(file_source)
            if not path.exists():
                raise FileAccessError(f"File not found: {path}")
            if not path.is_file():
                raise FileAccessError(f"Specified path is not a regular file: {path}")
            try:
                return path.read_bytes()
            except PermissionError as e:
                raise FileAccessError(f"Permission denied accessing file: {path}") from e
            except Exception as e:
                raise FileAccessError(f"Error reading file '{path}': {str(e)}") from e

        raise FileAccessError(f"Unsupported file source type: {type(file_source).__name__}")

    @classmethod
    def extract_text(
        cls,
        file_source: Union[str, Path, BinaryIO, bytes],
        file_type: Optional[str] = None,
        file_name: Optional[str] = None,
        raise_on_empty: bool = True,
    ) -> ProcessedDocument:
        """
        Universal entry point to extract clean text from educational documents.
        """
        resolved_ext = cls._resolve_file_type(file_source, file_type, file_name)
        file_bytes = cls._read_bytes(file_source)

        if len(file_bytes) == 0:
            if raise_on_empty:
                raise EmptyDocumentError("The provided document file is 0 bytes (empty).")
            return ProcessedDocument(
                file_name=file_name or (str(file_source) if isinstance(file_source, (str, Path)) else None),
                file_type=resolved_ext.lstrip("."),
                total_pages=0,
                pages=[],
                raw_text="",
                metadata={"file_size_bytes": 0, "is_empty": True},
            )

        inferred_name = file_name
        if not inferred_name and isinstance(file_source, (str, Path)):
            inferred_name = Path(file_source).name

        if resolved_ext == ".pdf":
            doc = cls._process_pdf(file_bytes, inferred_name)
        elif resolved_ext == ".docx":
            doc = cls._process_docx(file_bytes, inferred_name)
        elif resolved_ext in {".txt", ".md"}:
            doc = cls._process_txt(file_bytes, inferred_name, file_type_label=resolved_ext.lstrip("."))
        else:
            raise UnsupportedFileTypeError(f"Unsupported extension: {resolved_ext}")

        is_text_empty = not doc.raw_text or doc.raw_text.strip() == ""
        if is_text_empty:
            doc.metadata["is_empty"] = True
            if raise_on_empty:
                raise EmptyDocumentError(
                    f"No extractable text found in document '{inferred_name or 'unknown'}'. "
                    "The file may contain only scanned images or blank pages."
                )

        return doc

    @classmethod
    def extract_and_chunk(
        cls,
        file_source: Union[str, Path, BinaryIO, bytes],
        file_type: Optional[str] = None,
        file_name: Optional[str] = None,
        raise_on_empty: bool = True,
        chunker: Optional[Any] = None,
        **chunker_kwargs,
    ) -> Any:
        """
        Extract clean text and immediately return semantically chunked output.
        Unified bridge between Milestone 1 (extraction) and Milestone 2 (chunking).
        """
        processed_doc = cls.extract_text(
            file_source=file_source,
            file_type=file_type,
            file_name=file_name,
            raise_on_empty=raise_on_empty,
        )
        return processed_doc.chunk(chunker=chunker, **chunker_kwargs)

    @classmethod
    def _process_pdf(cls, file_bytes: bytes, file_name: Optional[str]) -> ProcessedDocument:
        """Extracts text page-by-page from a PDF document using pypdf."""
        try:
            stream = io.BytesIO(file_bytes)
            reader = pypdf.PdfReader(stream)
        except (pypdf.errors.PdfReadError, Exception) as e:
            raise CorruptedFileError(f"Failed to parse PDF document. The file may be corrupt: {str(e)}") from e

        if reader.is_encrypted:
            try:
                decrypted = reader.decrypt("")
                if decrypted == 0:
                    raise CorruptedFileError("PDF is password protected and cannot be read.")
            except Exception as e:
                raise CorruptedFileError(f"Failed to decrypt password-protected PDF: {str(e)}") from e

        pages: List[DocumentPage] = []
        raw_text_parts: List[str] = []

        try:
            total_pdf_pages = len(reader.pages)
            for idx, page in enumerate(reader.pages, start=1):
                try:
                    page_text = page.extract_text() or ""
                except Exception:
                    page_text = ""

                cleaned_page_text = cls.clean_text(page_text)
                words = cleaned_page_text.split() if cleaned_page_text else []
                
                doc_page = DocumentPage(
                    page_number=idx,
                    text=cleaned_page_text,
                    metadata={
                        "word_count": len(words),
                        "char_count": len(cleaned_page_text),
                    },
                )
                pages.append(doc_page)
                if cleaned_page_text:
                    raw_text_parts.append(cleaned_page_text)
        except Exception as e:
            raise CorruptedFileError(f"Error during PDF text extraction: {str(e)}") from e

        aggregated_text = "\n\n".join(raw_text_parts)
        total_words = sum(p.metadata.get("word_count", 0) for p in pages)
        total_chars = len(aggregated_text)

        doc_metadata = {
            "file_size_bytes": len(file_bytes),
            "word_count": total_words,
            "char_count": total_chars,
            "processed_at": datetime.now(timezone.utc).isoformat(),
        }

        if reader.metadata:
            try:
                if reader.metadata.title:
                    doc_metadata["title"] = str(reader.metadata.title)
                if reader.metadata.author:
                    doc_metadata["author"] = str(reader.metadata.author)
            except Exception:
                pass

        return ProcessedDocument(
            file_name=file_name,
            file_type="pdf",
            total_pages=len(pages),
            pages=pages,
            raw_text=aggregated_text,
            metadata=doc_metadata,
        )

    @classmethod
    def _process_docx(cls, file_bytes: bytes, file_name: Optional[str]) -> ProcessedDocument:
        """Extracts text from a DOCX document including paragraphs and tables."""
        try:
            stream = io.BytesIO(file_bytes)
            document = docx.Document(stream)
        except (docx.opc.exceptions.PackageNotFoundError, zipfile.BadZipFile, Exception) as e:
            raise CorruptedFileError(f"Failed to parse DOCX document. File is corrupted or invalid: {str(e)}") from e

        extracted_sections: List[str] = []

        try:
            for paragraph in document.paragraphs:
                p_text = paragraph.text.strip()
                if p_text:
                    extracted_sections.append(p_text)

            for table in document.tables:
                table_lines: List[str] = []
                for row in table.rows:
                    seen_cells = set()
                    row_cells: List[str] = []
                    for cell in row.cells:
                        if cell._tc not in seen_cells:
                            seen_cells.add(cell._tc)
                            cell_str = cell.text.strip().replace("\n", " ")
                            row_cells.append(cell_str)

                    if any(row_cells):
                        table_lines.append(" | ".join(row_cells))
                if table_lines:
                    extracted_sections.append("\n".join(table_lines))

        except Exception as e:
            raise CorruptedFileError(f"Error reading contents of DOCX document: {str(e)}") from e

        full_content = "\n\n".join(extracted_sections)
        cleaned_content = cls.clean_text(full_content)

        words = cleaned_content.split() if cleaned_content else []
        page = DocumentPage(
            page_number=1,
            text=cleaned_content,
            metadata={
                "word_count": len(words),
                "char_count": len(cleaned_content),
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
            },
        )

        doc_metadata = {
            "file_size_bytes": len(file_bytes),
            "word_count": len(words),
            "char_count": len(cleaned_content),
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "processed_at": datetime.now(timezone.utc).isoformat(),
        }

        try:
            core_props = document.core_properties
            if core_props.title:
                doc_metadata["title"] = core_props.title
            if core_props.author:
                doc_metadata["author"] = core_props.author
        except Exception:
            pass

        return ProcessedDocument(
            file_name=file_name,
            file_type="docx",
            total_pages=1,
            pages=[page],
            raw_text=cleaned_content,
            metadata=doc_metadata,
        )

    @classmethod
    def _process_txt(
        cls,
        file_bytes: bytes,
        file_name: Optional[str],
        file_type_label: str = "txt",
    ) -> ProcessedDocument:
        """Extracts text from plain text or markdown files with resilient encoding detection."""
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"]
        decoded_text: Optional[str] = None
        used_encoding: Optional[str] = None

        for enc in encodings:
            try:
                decoded_text = file_bytes.decode(enc)
                used_encoding = enc
                break
            except (UnicodeDecodeError, LookupError):
                continue

        if decoded_text is None:
            decoded_text = file_bytes.decode("utf-8", errors="replace")
            used_encoding = "utf-8 (replace)"

        raw_pages = decoded_text.split("\f")
        pages: List[DocumentPage] = []
        raw_text_parts: List[str] = []

        for idx, p_text in enumerate(raw_pages, start=1):
            cleaned = cls.clean_text(p_text)
            words = cleaned.split() if cleaned else []
            doc_page = DocumentPage(
                page_number=idx,
                text=cleaned,
                metadata={
                    "word_count": len(words),
                    "char_count": len(cleaned),
                },
            )
            pages.append(doc_page)
            if cleaned:
                raw_text_parts.append(cleaned)

        aggregated_text = "\n\n".join(raw_text_parts)
        total_words = sum(p.metadata.get("word_count", 0) for p in pages)
        total_chars = len(aggregated_text)

        doc_metadata = {
            "file_size_bytes": len(file_bytes),
            "encoding": used_encoding,
            "word_count": total_words,
            "char_count": total_chars,
            "processed_at": datetime.now(timezone.utc).isoformat(),
        }

        return ProcessedDocument(
            file_name=file_name,
            file_type=file_type_label,
            total_pages=len(pages),
            pages=pages,
            raw_text=aggregated_text,
            metadata=doc_metadata,
        )