"""
Comprehensive test suite for the DocumentProcessor module.
Covers PDF, DOCX, TXT extraction, stream inputs, error conditions, and serialization.
"""

import io
import tempfile
from pathlib import Path
import pytest
import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from app.rag.document_processor import (
    DocumentPage,
    DocumentProcessor,
    ProcessedDocument,
)
from app.rag.exceptions import (
    CorruptedFileError,
    DocumentProcessingError,
    EmptyDocumentError,
    FileAccessError,
    UnsupportedFileTypeError,
)


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """Generates a valid 2-page educational PDF as bytes."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(72, 750, "Introduction to Physics: Mechanics")
    c.drawString(72, 720, "Newton's First Law of Motion: An object remains at rest unless acted upon.")
    c.showPage()
    c.drawString(72, 750, "Chapter 2: Energy and Work")
    c.drawString(72, 720, "Work is defined as force multiplied by displacement in the direction of force.")
    c.showPage()
    c.save()
    return buf.getvalue()


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """Generates a valid DOCX document with headings, paragraphs, and tables as bytes."""
    doc = docx.Document()
    doc.add_heading("Computer Networks 101", level=1)
    doc.add_paragraph("The OSI model characterizes and standardizes communication functions.")
    
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Layer"
    table.rows[0].cells[1].text = "Protocol"
    table.rows[1].cells[1].text = "HTTP, DNS"
    table.rows[1].cells[0].text = "Application"
    table.rows[2].cells[0].text = "Transport"
    table.rows[2].cells[1].text = "TCP, UDP"
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ----------------------------------------------------------------------
# PDF Extraction Tests
# ----------------------------------------------------------------------

def test_extract_pdf_from_bytes(sample_pdf_bytes):
    """Test extracting text and page structure from PDF in-memory bytes."""
    result = DocumentProcessor.extract_text(
        file_source=sample_pdf_bytes,
        file_type="pdf",
        file_name="physics_notes.pdf",
    )

    assert isinstance(result, ProcessedDocument)
    assert result.file_name == "physics_notes.pdf"
    assert result.file_type == "pdf"
    assert result.total_pages == 2
    assert len(result.pages) == 2

    # Page 1 checks
    assert result.pages[0].page_number == 1
    assert "Introduction to Physics: Mechanics" in result.pages[0].text
    assert "Newton's First Law" in result.pages[0].text
    assert result.pages[0].metadata["word_count"] > 0

    # Page 2 checks
    assert result.pages[1].page_number == 2
    assert "Chapter 2: Energy and Work" in result.pages[1].text

    # Aggregated raw text checks
    assert "Introduction to Physics" in result.raw_text
    assert "Chapter 2: Energy" in result.raw_text
    assert result.metadata["word_count"] > 0
    assert result.metadata["file_size_bytes"] == len(sample_pdf_bytes)


def test_extract_pdf_from_file_path(tmp_path, sample_pdf_bytes):
    """Test extracting text from a PDF file path on disk."""
    pdf_path = tmp_path / "lecture1.pdf"
    pdf_path.write_bytes(sample_pdf_bytes)

    result = DocumentProcessor.extract_text(pdf_path)
    assert result.file_name == "lecture1.pdf"
    assert result.file_type == "pdf"
    assert result.total_pages == 2
    assert "Newton's First Law" in result.raw_text


def test_extract_pdf_from_stream(sample_pdf_bytes):
    """Test extracting text from a file-like BytesIO stream."""
    stream = io.BytesIO(sample_pdf_bytes)
    result = DocumentProcessor.extract_text(stream, file_type="application/pdf", file_name="stream_file.pdf")
    assert result.total_pages == 2
    assert "Introduction to Physics" in result.pages[0].text


# ----------------------------------------------------------------------
# DOCX Extraction Tests
# ----------------------------------------------------------------------

def test_extract_docx_from_bytes(sample_docx_bytes):
    """Test extracting text, paragraphs, and tables from DOCX in-memory bytes."""
    result = DocumentProcessor.extract_text(
        file_source=sample_docx_bytes,
        file_type="docx",
        file_name="networking.docx",
    )

    assert result.file_name == "networking.docx"
    assert result.file_type == "docx"
    assert result.total_pages == 1
    assert "Computer Networks 101" in result.raw_text
    assert "OSI model" in result.raw_text
    # Verify table row extraction
    assert "Layer | Protocol" in result.raw_text
    assert "Application | HTTP, DNS" in result.raw_text
    assert "Transport | TCP, UDP" in result.raw_text
    assert result.metadata["table_count"] == 1


def test_extract_docx_from_file_path(tmp_path, sample_docx_bytes):
    """Test extracting text from a DOCX file path on disk."""
    docx_path = tmp_path / "guide.docx"
    docx_path.write_bytes(sample_docx_bytes)

    result = DocumentProcessor.extract_text(str(docx_path))
    assert result.file_name == "guide.docx"
    assert result.file_type == "docx"
    assert "OSI model" in result.raw_text


# ----------------------------------------------------------------------
# TXT & Markdown Extraction Tests
# ----------------------------------------------------------------------

def test_extract_txt_basic(tmp_path):
    """Test basic plain text file extraction."""
    txt_content = "Linear Algebra Lecture\n\nVectors: Elements of vector spaces.\nMatrices: Linear maps."
    txt_file = tmp_path / "math.txt"
    txt_file.write_text(txt_content, encoding="utf-8")

    result = DocumentProcessor.extract_text(txt_file)
    assert result.file_name == "math.txt"
    assert result.file_type == "txt"
    assert result.total_pages == 1
    assert "Vectors: Elements of vector spaces." in result.raw_text
    assert result.metadata["word_count"] == 11


def test_extract_txt_with_form_feed_pagination():
    """Test form-feed character (\f) pagination support in TXT files."""
    content = "Section 1: Foundations\fSection 2: Intermediate\fSection 3: Advanced"
    result = DocumentProcessor.extract_text(
        content.encode("utf-8"),
        file_type="txt",
        file_name="paginated_notes.txt",
    )

    assert result.total_pages == 3
    assert result.pages[0].page_number == 1
    assert result.pages[0].text == "Section 1: Foundations"
    assert result.pages[1].page_number == 2
    assert result.pages[1].text == "Section 2: Intermediate"
    assert result.pages[2].page_number == 3
    assert result.pages[2].text == "Section 3: Advanced"


def test_extract_txt_with_encodings():
    """Test handling of different encodings (Latin-1, UTF-8 with BOM)."""
    latin1_bytes = "Cafe resume naive Schrodinger".encode("latin-1")
    result_latin1 = DocumentProcessor.extract_text(latin1_bytes, file_type="txt")
    assert "Cafe" in result_latin1.raw_text
    assert "Schrodinger" in result_latin1.raw_text

    utf8_bom_bytes = "Machine Learning Course".encode("utf-8-sig")
    result_bom = DocumentProcessor.extract_text(utf8_bom_bytes, file_type="txt")
    assert "Machine Learning Course" in result_bom.raw_text


def test_extract_markdown():
    """Test markdown file extraction."""
    md_content = "# Title\n\n- Point 1\n- Point 2"
    result = DocumentProcessor.extract_text(md_content.encode("utf-8"), file_name="notes.md")
    assert result.file_type == "md"
    assert "Point 1" in result.raw_text


# ----------------------------------------------------------------------
# Text Cleaning & Normalization Tests
# ----------------------------------------------------------------------

def test_clean_text_normalization():
    """Test text cleaning, Unicode normalization, control char removal, and line break cleanup."""
    raw = "Line 1   \r\nLine 2\x00\x07\n\n\n\nLine 3 with tabs\tpreserved"
    cleaned = DocumentProcessor.clean_text(raw)
    
    assert "\x00" not in cleaned
    assert "\x07" not in cleaned
    assert "Line 1\nLine 2\n\nLine 3 with tabs\tpreserved" == cleaned


# ----------------------------------------------------------------------
# Error Handling & Edge Cases
# ----------------------------------------------------------------------

def test_empty_file_raises_error():
    """Test that empty 0-byte input raises EmptyDocumentError when raise_on_empty=True."""
    with pytest.raises(EmptyDocumentError, match="0 bytes"):
        DocumentProcessor.extract_text(b"", file_type="txt")


def test_empty_file_returns_empty_doc_when_not_raising():
    """Test that empty 0-byte input returns clean empty doc when raise_on_empty=False."""
    result = DocumentProcessor.extract_text(b"", file_type="txt", raise_on_empty=False)
    assert result.total_pages == 0
    assert result.raw_text == ""
    assert result.metadata["is_empty"] is True


def test_blank_whitespace_file_raises_error():
    """Test that whitespace-only file raises EmptyDocumentError."""
    with pytest.raises(EmptyDocumentError, match="No extractable text"):
        DocumentProcessor.extract_text(b"    \n\n   \t  ", file_type="txt")


def test_corrupted_pdf_raises_corrupted_file_error():
    """Test that invalid/corrupted PDF bytes raise CorruptedFileError."""
    corrupted_pdf = b"%PDF-1.4\nCorrupted binary payload that is completely invalid"
    with pytest.raises(CorruptedFileError):
        DocumentProcessor.extract_text(corrupted_pdf, file_type="pdf")


def test_corrupted_docx_raises_corrupted_file_error():
    """Test that invalid/non-zip DOCX bytes raise CorruptedFileError."""
    corrupted_docx = b"PK\x03\x04not a valid word document zip file"
    with pytest.raises(CorruptedFileError):
        DocumentProcessor.extract_text(corrupted_docx, file_type="docx")


def test_unsupported_file_type_raises_error():
    """Test that unsupported file extensions raise UnsupportedFileTypeError."""
    with pytest.raises(UnsupportedFileTypeError, match="Unsupported or indeterminate file type"):
        DocumentProcessor.extract_text(b"some data", file_name="image.png")

    with pytest.raises(UnsupportedFileTypeError):
        DocumentProcessor.extract_text(b"some data", file_type="application/zip")


def test_missing_file_raises_file_access_error():
    """Test that non-existent file path raises FileAccessError."""
    with pytest.raises(FileAccessError, match="File not found"):
        DocumentProcessor.extract_text(Path("non_existent_directory/missing_doc.pdf"))


# ----------------------------------------------------------------------
# Serialization Tests
# ----------------------------------------------------------------------

def test_to_dict_serialization(sample_pdf_bytes):
    """Test JSON/dict serialization of ProcessedDocument and DocumentPage."""
    result = DocumentProcessor.extract_text(
        sample_pdf_bytes,
        file_name="sample.pdf",
        file_type="pdf",
    )

    doc_dict = result.to_dict()
    assert isinstance(doc_dict, dict)
    assert doc_dict["file_name"] == "sample.pdf"
    assert doc_dict["file_type"] == "pdf"
    assert doc_dict["total_pages"] == 2
    assert isinstance(doc_dict["pages"], list)
    assert len(doc_dict["pages"]) == 2
    
    first_page = doc_dict["pages"][0]
    assert first_page["page_number"] == 1
    assert "Introduction to Physics" in first_page["text"]
    assert "word_count" in first_page["metadata"]
    assert "char_count" in first_page["metadata"]


# ----------------------------------------------------------------------
# Extended Stream & Edge Cases Tests
# ----------------------------------------------------------------------

def test_spooled_temporary_file_stream():
    """Test extracting from a SpooledTemporaryFile (FastAPI UploadFile structure)."""
    with tempfile.SpooledTemporaryFile(max_size=1024 * 1024, mode="w+b") as temp_file:
        temp_file.write(b"Course Title: Artificial Intelligence\nModule 1: Search algorithms.")
        temp_file.seek(0)
        
        result = DocumentProcessor.extract_text(temp_file, file_type="txt", file_name="ai_syllabus.txt")
        assert result.total_pages == 1
        assert "Artificial Intelligence" in result.raw_text
        assert "Search algorithms" in result.raw_text


def test_docx_multiple_tables_and_mixed_content():
    """Test DOCX with multiple tables, empty paragraphs, and mixed formatting."""
    doc = docx.Document()
    doc.add_heading("Course Schedule", level=1)
    doc.add_paragraph("")
    
    table1 = doc.add_table(rows=2, cols=2)
    table1.rows[0].cells[0].text = "Week"
    table1.rows[0].cells[1].text = "Topic"
    table1.rows[1].cells[0].text = "1"
    table1.rows[1].cells[1].text = "Probability"
    
    doc.add_paragraph("Midterm Exam will take place in Week 5.")
    
    table2 = doc.add_table(rows=2, cols=2)
    table2.rows[0].cells[0].text = "Assignment"
    table2.rows[0].cells[1].text = "Weight"
    table2.rows[1].cells[0].text = "HW 1"
    table2.rows[1].cells[1].text = "15%"
    
    buf = io.BytesIO()
    doc.save(buf)
    
    result = DocumentProcessor.extract_text(buf.getvalue(), file_name="syllabus.docx")
    assert result.total_pages == 1
    assert "Week | Topic" in result.raw_text
    assert "Probability" in result.raw_text
    assert "Midterm Exam" in result.raw_text
    assert "Assignment | Weight" in result.raw_text
    assert result.metadata["table_count"] == 2


def test_unicode_and_math_symbols():
    """Test extraction of Greek mathematical symbols, equations, and Unicode text."""
    math_text = "Calculus II\nIntegral: \u222b f(x)dx\nSum: \u2211_{i=1}^n x_i\nAlpha & Beta: \u03b1 + \u03b2 = \u03b3"
    result = DocumentProcessor.extract_text(math_text.encode("utf-8"), file_name="math.txt")
    assert "\u222b" in result.raw_text
    assert "\u2211" in result.raw_text
    assert "\u03b1" in result.raw_text